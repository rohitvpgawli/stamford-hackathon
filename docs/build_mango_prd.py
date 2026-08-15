from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parent
SOURCE = ROOT / "Mango_Hackathon_PRD.md"
OUTPUT = ROOT / "Mango_Hackathon_PRD.docx"

# compact_reference_guide preset, with one named brand override:
# cover title and select emphasis use Mango orange (#C96A11).
PAGE_WIDTH_DXA = 12240
CONTENT_WIDTH_DXA = 9360
TABLE_INDENT_DXA = 120
CELL_MARGINS_DXA = {"top": 80, "bottom": 80, "start": 120, "end": 120}

INK = "172B4D"
HEADING_BLUE = "2E74B5"
HEADING_DARK = "1F4D78"
MANGO = "C96A11"
MANGO_LIGHT = "FFF3E5"
BLUE_GRAY = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
MUTED = "5F6B7A"
WHITE = "FFFFFF"
BORDER = "CBD5E1"
CODE_FILL = "F7F8FA"


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for key, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{key}"))
        if node is None:
            node = OxmlElement(f"w:{key}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_width(cell, width_dxa):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths):
    assert sum(widths) == CONTENT_WIDTH_DXA
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl_pr = table._tbl.tblPr

    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(CONTENT_WIDTH_DXA))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(TABLE_INDENT_DXA))
    tbl_ind.set(qn("w:type"), "dxa")

    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            set_cell_width(cell, widths[idx])
            set_cell_margins(cell, **CELL_MARGINS_DXA)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_borders(cell, color=BORDER, size="4"):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = tc_pr.first_child_found_in("w:tcBorders")
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)
    for edge in ("top", "start", "bottom", "end", "insideH", "insideV"):
        tag = tc_borders.find(qn(f"w:{edge}"))
        if tag is None:
            tag = OxmlElement(f"w:{edge}")
            tc_borders.append(tag)
        tag.set(qn("w:val"), "single")
        tag.set(qn("w:sz"), size)
        tag.set(qn("w:color"), color)


def set_paragraph_border(paragraph, side="left", color=MANGO, size="18", space="8"):
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    border = OxmlElement(f"w:{side}")
    border.set(qn("w:val"), "single")
    border.set(qn("w:sz"), size)
    border.set(qn("w:space"), space)
    border.set(qn("w:color"), color)
    p_bdr.append(border)


def set_paragraph_shading(paragraph, fill):
    p_pr = paragraph._p.get_or_add_pPr()
    shd = p_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        p_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def add_field(paragraph, instruction):
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "Update field in Word if needed"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run = paragraph.add_run()._r
    run.append(begin)
    run.append(instr)
    run.append(separate)
    run.append(text)
    run.append(end)


def add_page_number(paragraph):
    run = paragraph.add_run("Page ")
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor.from_string(MUTED)
    add_field(paragraph, " PAGE ")


def set_font(run, name="Calibri", size=None, color=None, bold=None, italic=None):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def apply_inline(paragraph, text, default_size=11, default_color=INK):
    token_re = re.compile(r"(\*\*.+?\*\*|`.+?`|\*[^*]+?\*)")
    pos = 0
    for match in token_re.finditer(text):
        if match.start() > pos:
            run = paragraph.add_run(text[pos:match.start()])
            set_font(run, size=default_size, color=default_color)
        token = match.group(0)
        if token.startswith("**"):
            run = paragraph.add_run(token[2:-2])
            set_font(run, size=default_size, color=default_color, bold=True)
        elif token.startswith("`"):
            run = paragraph.add_run(token[1:-1])
            set_font(run, name="Consolas", size=max(default_size - 1, 8.5), color=HEADING_DARK)
            run.font.highlight_color = None
        else:
            run = paragraph.add_run(token[1:-1])
            set_font(run, size=default_size, color=default_color, italic=True)
        pos = match.end()
    if pos < len(text):
        run = paragraph.add_run(text[pos:])
        set_font(run, size=default_size, color=default_color)


def create_numbering(doc):
    numbering = doc.part.numbering_part.element

    def add_definition(abstract_id, num_id, marker, fmt, left, hanging, font="Calibri"):
        abstract = OxmlElement("w:abstractNum")
        abstract.set(qn("w:abstractNumId"), str(abstract_id))
        multi = OxmlElement("w:multiLevelType")
        multi.set(qn("w:val"), "singleLevel")
        abstract.append(multi)
        lvl = OxmlElement("w:lvl")
        lvl.set(qn("w:ilvl"), "0")
        start = OxmlElement("w:start")
        start.set(qn("w:val"), "1")
        lvl.append(start)
        num_fmt = OxmlElement("w:numFmt")
        num_fmt.set(qn("w:val"), fmt)
        lvl.append(num_fmt)
        lvl_text = OxmlElement("w:lvlText")
        lvl_text.set(qn("w:val"), marker)
        lvl.append(lvl_text)
        suff = OxmlElement("w:suff")
        suff.set(qn("w:val"), "tab")
        lvl.append(suff)
        p_pr = OxmlElement("w:pPr")
        tabs = OxmlElement("w:tabs")
        tab = OxmlElement("w:tab")
        tab.set(qn("w:val"), "num")
        tab.set(qn("w:pos"), str(left))
        tabs.append(tab)
        p_pr.append(tabs)
        ind = OxmlElement("w:ind")
        ind.set(qn("w:left"), str(left))
        ind.set(qn("w:hanging"), str(hanging))
        p_pr.append(ind)
        lvl.append(p_pr)
        r_pr = OxmlElement("w:rPr")
        fonts = OxmlElement("w:rFonts")
        fonts.set(qn("w:ascii"), font)
        fonts.set(qn("w:hAnsi"), font)
        r_pr.append(fonts)
        lvl.append(r_pr)
        abstract.append(lvl)
        numbering.append(abstract)

        num = OxmlElement("w:num")
        num.set(qn("w:numId"), str(num_id))
        abstract_ref = OxmlElement("w:abstractNumId")
        abstract_ref.set(qn("w:val"), str(abstract_id))
        num.append(abstract_ref)
        numbering.append(num)

    add_definition(90, 90, "•", "bullet", 540, 270, "Symbol")
    add_definition(91, 91, "%1.", "decimal", 540, 270)
    return 90, 91


def create_number_instance(doc, num_id, abstract_id=91):
    numbering = doc.part.numbering_part.element
    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    override = OxmlElement("w:lvlOverride")
    override.set(qn("w:ilvl"), "0")
    start_override = OxmlElement("w:startOverride")
    start_override.set(qn("w:val"), "1")
    override.append(start_override)
    num.append(override)
    numbering.append(num)


def attach_numbering(paragraph, num_id):
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num_id_el = OxmlElement("w:numId")
    num_id_el.set(qn("w:val"), str(num_id))
    num_pr.append(ilvl)
    num_pr.append(num_id_el)
    p_pr.append(num_pr)


def choose_widths(rows):
    cols = len(rows[0])
    if cols == 1:
        return [CONTENT_WIDTH_DXA]
    max_lengths = [max(len(re.sub(r"[*`]", "", row[i])) for row in rows) for i in range(cols)]
    floors = [0.13] * cols
    if cols == 2:
        floors = [0.25, 0.55]
    weights = [max(floors[i], min(0.5, max_lengths[i] / max(sum(max_lengths), 1))) for i in range(cols)]
    total = sum(weights)
    raw = [max(1100, round(CONTENT_WIDTH_DXA * w / total)) for w in weights]
    scale = CONTENT_WIDTH_DXA / sum(raw)
    widths = [round(v * scale) for v in raw]
    widths[-1] += CONTENT_WIDTH_DXA - sum(widths)
    return widths


def add_table(doc, rows):
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.style = "Table Grid"
    widths = choose_widths(rows)
    set_table_geometry(table, widths)
    for r_idx, row in enumerate(rows):
        for c_idx, value in enumerate(row):
            cell = table.cell(r_idx, c_idx)
            set_cell_borders(cell)
            if r_idx == 0:
                shade_cell(cell, BLUE_GRAY)
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(1)
            p.paragraph_format.line_spacing = 1.05
            apply_inline(p, value.strip(), default_size=9.2, default_color=INK)
            if r_idx == 0:
                for run in p.runs:
                    run.bold = True
                    run.font.color.rgb = RGBColor.from_string(HEADING_DARK)
    set_repeat_table_header(table.rows[0])
    after = doc.add_paragraph()
    after.paragraph_format.space_after = Pt(2)


def setup_document():
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    section.header_distance = Inches(0.42)
    section.footer_distance = Inches(0.42)
    section.different_first_page_header_footer = False

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.15

    for style_name, size, color, before, after in (
        ("Title", 27, MANGO, 0, 10),
        ("Subtitle", 13, MUTED, 0, 18),
        ("Heading 1", 16, HEADING_BLUE, 18, 10),
        ("Heading 2", 13, HEADING_BLUE, 14, 7),
        ("Heading 3", 12, HEADING_DARK, 10, 5),
    ):
        style = styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = style_name != "Subtitle"
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    # Word's built-in Title style can carry a decorative paragraph border.
    # The chosen memo masthead uses typography, not a title underline.
    title_ppr = styles["Title"]._element.get_or_add_pPr()
    title_border = title_ppr.find(qn("w:pBdr"))
    if title_border is not None:
        title_ppr.remove(title_border)

    if "Code Block" not in styles:
        code = styles.add_style("Code Block", WD_STYLE_TYPE.PARAGRAPH)
    else:
        code = styles["Code Block"]
    code.font.name = "Consolas"
    code._element.rPr.rFonts.set(qn("w:ascii"), "Consolas")
    code._element.rPr.rFonts.set(qn("w:hAnsi"), "Consolas")
    code.font.size = Pt(8.3)
    code.font.color.rgb = RGBColor.from_string(HEADING_DARK)
    code.paragraph_format.left_indent = Inches(0.16)
    code.paragraph_format.right_indent = Inches(0.08)
    code.paragraph_format.space_before = Pt(2)
    code.paragraph_format.space_after = Pt(2)
    code.paragraph_format.line_spacing = 1.0

    bullet_num_id, decimal_num_id = create_numbering(doc)
    return doc, bullet_num_id, decimal_num_id


def add_running_furniture(section):
    # A restrained footer is more robust across Word and LibreOffice than
    # alternating running headers, while the heading ladder provides navigation.
    for footer in (section.footer, section.even_page_footer):
        fp = footer.paragraphs[0]
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        fp.paragraph_format.space_before = Pt(0)
        add_page_number(fp)


def add_cover(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(20)
    run = p.add_run("PRODUCT REQUIREMENTS DOCUMENT")
    set_font(run, size=10, color=MANGO, bold=True)

    title = doc.add_paragraph(style="Title")
    title.add_run("Mango Hackathon Build")
    title.paragraph_format.space_after = Pt(3)

    subtitle = doc.add_paragraph(style="Subtitle")
    subtitle.add_run("Android local-SIM SMS gateway • Hermes Agent on EC2 • Personalized app mockup")
    subtitle.paragraph_format.space_after = Pt(18)

    metadata = [
        ("Status", "Implementation-ready v1.0"),
        ("Date", "August 15, 2026"),
        ("Build posture", "Deliberately scrappy hackathon prototype"),
        ("Target load", "50–100 unique participants"),
        ("Owner", "Mango / Trillium"),
    ]
    for label, value in metadata:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run(f"{label}: ")
        set_font(r, size=10.5, color=INK, bold=True)
        r = p.add_run(value)
        set_font(r, size=10.5, color=INK)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(24)
    p.paragraph_format.space_after = Pt(12)
    set_paragraph_shading(p, MANGO_LIGHT)
    set_paragraph_border(p, "left", MANGO, "24", "10")
    r = p.add_run("Product thesis\n")
    set_font(r, size=9.5, color=MANGO, bold=True)
    r = p.add_run("Text for recommendation. App for exploration.")
    set_font(r, size=16, color=INK, bold=True)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run("Implementation directive")
    set_font(r, size=11, color=HEADING_DARK, bold=True)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    apply_inline(
        p,
        "Build one reliable end-to-end story first: real SMS round trip → isolated session → grounded recommendation → JOIN → personalized app handoff.",
        default_size=11,
    )

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    r = p.add_run("CONFIDENTIALITY / DATA NOTE")
    set_font(r, size=8.5, color=MUTED, bold=True)
    p = doc.add_paragraph()
    apply_inline(
        p,
        "Hackathon-only prototype. Seeded events, places, and people must be labeled as demo data; do not present them as verified live listings.",
        default_size=9.5,
        default_color=MUTED,
    )


def add_toc(doc):
    h = doc.add_paragraph("Contents", style="Heading 1")
    h.paragraph_format.space_before = Pt(0)
    intro = doc.add_paragraph()
    apply_inline(intro, "This PRD is organized as an operator’s reference. Use Word’s Navigation Pane for direct section links.", 10.5, MUTED)

    groups = [
        ("PRODUCT + EXPERIENCE", "0–6", "Executive summary, goals, users, principles, MVP scope, happy flows, and experience requirements."),
        ("SYSTEM + STATE", "7–11", "Architecture, data model, intent routing, session/turn policy, and recommendation/matching logic."),
        ("AGENT + SAFETY + CONTRACTS", "12–20", "Hermes behavior, guardrails, APIs, concurrency, app mockup, seed data, errors, privacy, and operations."),
        ("SHIP + VERIFY", "21–29", "Acceptance criteria, tests, build priorities, module layout, risks, demo script, metrics, decisions, and Codex directive."),
    ]
    for label, section_range, description in groups:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run(f"{label}  ·  Sections {section_range}")
        set_font(r, size=10, color=MANGO, bold=True)
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.16)
        p.paragraph_format.space_after = Pt(5)
        apply_inline(p, description, 10.5, INK)

    h2 = doc.add_paragraph("Defaults to freeze before coding", style="Heading 2")
    h2.paragraph_format.space_before = Pt(14)
    defaults = [
        "One Android local-SIM gateway behind a vendor-neutral adapter.",
        "One EC2 TypeScript service; SQLite WAL unless Postgres already exists.",
        "12-hour sessions; warning at turn 10; cutoff at turn 12.",
        "One recommendation, one alternative, then the personalized app handoff.",
        "JOIN may finish through two protected turns beyond the normal budget.",
        "Seeded 18+ people and clearly labeled demo opportunities.",
    ]
    for item in defaults:
        p = doc.add_paragraph()
        attach_numbering(p, 90)
        p.paragraph_format.space_after = Pt(3)
        apply_inline(p, item, 10.2, INK)
    doc.add_page_break()


def parse_markdown(doc, bullet_num_id, decimal_num_id):
    lines = SOURCE.read_text(encoding="utf-8").splitlines()
    idx = 0
    in_code = False
    code_buffer = []
    last_block_type = None
    next_decimal_num_id = 100
    current_decimal_num_id = decimal_num_id

    while idx < len(lines):
        line = lines[idx]
        stripped = line.strip()

        if stripped.startswith("```"):
            if not in_code:
                in_code = True
                code_buffer = []
            else:
                for code_line in code_buffer or [""]:
                    p = doc.add_paragraph(style="Code Block")
                    set_paragraph_shading(p, CODE_FILL)
                    r = p.add_run(code_line if code_line else " ")
                    set_font(r, name="Consolas", size=8.3, color=HEADING_DARK)
                in_code = False
                code_buffer = []
            idx += 1
            continue

        if in_code:
            code_buffer.append(line)
            idx += 1
            continue

        if not stripped:
            last_block_type = None
            idx += 1
            continue

        if stripped.startswith("# "):
            # Title and metadata are represented on the designed cover.
            idx += 1
            while idx < len(lines) and (not lines[idx].strip() or lines[idx].startswith("**")):
                idx += 1
            continue

        heading_match = re.match(r"^(#{2,4})\s+(.+)$", stripped)
        if heading_match:
            hashes, heading = heading_match.groups()
            level = len(hashes) - 1
            p = doc.add_paragraph(style=f"Heading {min(level, 3)}")
            p.add_run(heading)
            last_block_type = "heading"
            idx += 1
            continue

        if stripped.startswith("|") and idx + 1 < len(lines) and re.match(r"^\s*\|?\s*:?-+", lines[idx + 1]):
            raw_rows = []
            while idx < len(lines) and lines[idx].strip().startswith("|"):
                raw_rows.append(lines[idx].strip())
                idx += 1
            parsed = []
            for ridx, row in enumerate(raw_rows):
                cells = [cell.strip() for cell in row.strip("|").split("|")]
                if ridx == 1 and all(re.match(r"^:?-+:?$", cell) for cell in cells):
                    continue
                parsed.append(cells)
            if parsed and all(len(row) == len(parsed[0]) for row in parsed):
                add_table(doc, parsed)
            last_block_type = "table"
            continue

        if stripped.startswith("> "):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.18)
            p.paragraph_format.right_indent = Inches(0.08)
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(8)
            set_paragraph_shading(p, MANGO_LIGHT)
            set_paragraph_border(p, "left", MANGO, "18", "8")
            apply_inline(p, stripped[2:], 10.5, INK)
            last_block_type = "quote"
            idx += 1
            continue

        checklist = re.match(r"^- \[ \]\s+(.+)$", stripped)
        if checklist:
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.12)
            p.paragraph_format.first_line_indent = Inches(-0.02)
            p.paragraph_format.space_after = Pt(4)
            r = p.add_run("☐  ")
            set_font(r, name="Arial", size=10.5, color=MANGO)
            apply_inline(p, checklist.group(1), 10.5, INK)
            last_block_type = "checklist"
            idx += 1
            continue

        bullet = re.match(r"^-\s+(.+)$", stripped)
        if bullet:
            p = doc.add_paragraph()
            attach_numbering(p, bullet_num_id)
            p.paragraph_format.space_after = Pt(4)
            p.paragraph_format.line_spacing = 1.15
            apply_inline(p, bullet.group(1), 10.7, INK)
            last_block_type = "bullet"
            idx += 1
            continue

        numbered = re.match(r"^\d+\.\s+(.+)$", stripped)
        if numbered:
            if last_block_type != "numbered":
                current_decimal_num_id = next_decimal_num_id
                create_number_instance(doc, current_decimal_num_id)
                next_decimal_num_id += 1
            p = doc.add_paragraph()
            attach_numbering(p, current_decimal_num_id)
            p.paragraph_format.space_after = Pt(4)
            p.paragraph_format.line_spacing = 1.15
            apply_inline(p, numbered.group(1), 10.7, INK)
            last_block_type = "numbered"
            idx += 1
            continue

        # Coalesce consecutive prose lines into a single paragraph.
        prose = [stripped]
        idx += 1
        while idx < len(lines):
            candidate = lines[idx].strip()
            if not candidate:
                break
            if (
                candidate.startswith("#")
                or candidate.startswith("|")
                or candidate.startswith("```")
                or candidate.startswith("> ")
                or re.match(r"^-\s+", candidate)
                or re.match(r"^\d+\.\s+", candidate)
            ):
                break
            prose.append(candidate)
            idx += 1
        p = doc.add_paragraph()
        apply_inline(p, " ".join(prose), 11, INK)
        last_block_type = "prose"


def audit_document(doc):
    section = doc.sections[0]
    emu_per_inch = 914400
    assert round(section.page_width / emu_per_inch, 2) == 8.5
    assert round(section.page_height / emu_per_inch, 2) == 11.0
    assert round(section.left_margin / emu_per_inch, 2) == 1.0
    assert round(section.right_margin / emu_per_inch, 2) == 1.0
    for table in doc.tables:
        grid = table._tbl.tblGrid
        widths = [int(col.get(qn("w:w"))) for col in grid]
        assert sum(widths) == CONTENT_WIDTH_DXA, widths
        assert all(width > 0 for width in widths)
        for row in table.rows:
            assert len(row.cells) == len(widths)
    for paragraph in doc.paragraphs:
        if paragraph.style.name.startswith("Heading"):
            assert paragraph.style.paragraph_format.keep_with_next


def main():
    doc, bullet_num_id, decimal_num_id = setup_document()
    add_cover(doc)
    body_section = doc.add_section(WD_SECTION.NEW_PAGE)
    body_section.page_width = Inches(8.5)
    body_section.page_height = Inches(11)
    body_section.top_margin = Inches(0.8)
    body_section.bottom_margin = Inches(0.75)
    body_section.left_margin = Inches(1.0)
    body_section.right_margin = Inches(1.0)
    body_section.header_distance = Inches(0.42)
    body_section.footer_distance = Inches(0.42)
    body_section.different_first_page_header_footer = False
    for part in (
        body_section.header,
        body_section.even_page_header,
        body_section.footer,
        body_section.even_page_footer,
    ):
        part.is_linked_to_previous = False
    add_toc(doc)
    parse_markdown(doc, bullet_num_id, decimal_num_id)

    core = doc.core_properties
    core.title = "Mango Hackathon Build — Product Requirements Document"
    core.subject = "Implementation-ready PRD for the Mango Android SMS + Hermes hackathon prototype"
    core.author = "Mango / Trillium"
    core.keywords = "Mango, Stamford, SMS, Android gateway, Hermes, hackathon, PRD"
    core.comments = "Generated from the implementation-ready Markdown source in the Mango workspace."

    settings = doc.settings._element
    even_odd = settings.find(qn("w:evenAndOddHeaders"))
    if even_odd is None:
        even_odd = OxmlElement("w:evenAndOddHeaders")
        settings.append(even_odd)
    even_odd.set(qn("w:val"), "true")
    update = settings.find(qn("w:updateFields"))
    if update is None:
        update = OxmlElement("w:updateFields")
        settings.append(update)
    update.set(qn("w:val"), "true")

    audit_document(doc)
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
